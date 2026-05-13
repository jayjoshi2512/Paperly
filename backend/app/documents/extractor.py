from dataclasses import dataclass
from typing import List
import fitz  # PyMuPDF
import io

@dataclass
class PageText:
    page_number: int
    text: str
    char_count: int

class ExtractionError(Exception):
    pass

def extract_text_from_pdf(file_bytes: bytes) -> List[PageText]:
    """
    Extracts text from PDF bytes.
    Strips headers and footers heuristically (top/bottom 5% of the page).
    """
    if not file_bytes:
        return []
        
    # Check magic number for PDF
    if not file_bytes.startswith(b'%PDF'):
        raise ExtractionError("Invalid file type: Not a PDF")
        
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        # Corrupt PDF
        return []
        
    pages = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Get dimensions
        rect = page.rect
        page_height = rect.height
        
        top_margin = page_height * 0.05
        bottom_margin = page_height * 0.95
        
        # Extract blocks: (x0, y0, x1, y1, "text", block_no, block_type)
        blocks = page.get_text("blocks")
        
        content_blocks = []
        for b in blocks:
            # Check if block is text (block_type == 0)
            if b[6] == 0:
                y0 = b[1]
                y1 = b[3]
                
                # Check if block is outside top/bottom 5%
                if y0 >= top_margin and y1 <= bottom_margin:
                    content_blocks.append(b[4].strip())
                    
        page_text = "\n".join(content_blocks).strip()
        if page_text:
            pages.append(PageText(
                page_number=page_num + 1,
                text=page_text,
                char_count=len(page_text)
            ))
            
    doc.close()
    return pages


def extract_text_from_docx(file_bytes: bytes) -> List[PageText]:
    """
    Extract and virtually paginate text from a DOCX file.

    DOCX has no concept of physical pages. We simulate virtual pages by
    grouping every 500 words into one PageText, which allows downstream
    citation code to say "virtual page 3" rather than "unknown".

    Tables are extracted as pipe-separated rows to preserve structure.
    Headings are preserved as plain text. Empty paragraphs are skipped.
    """
    try:
        from docx import Document as DocxDocument  # python-docx
    except ImportError as e:
        raise ExtractionError("python-docx is not installed: pip install python-docx") from e

    if not file_bytes:
        return []

    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise ExtractionError(f"Invalid or corrupt DOCX file: {e}") from e

    paragraphs: List[str] = []

    # Extract body paragraphs (includes headings — their text is preserved)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extract tables: each row becomes a pipe-separated line
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    if not paragraphs:
        return []

    # Group into virtual pages of ~500 words each
    pages: List[PageText] = []
    page_num = 1
    current_page: List[str] = []
    word_count = 0

    for para in paragraphs:
        para_words = len(para.split())
        current_page.append(para)
        word_count += para_words

        if word_count >= 500:
            page_text = "\n".join(current_page)
            pages.append(PageText(
                page_number=page_num,
                text=page_text,
                char_count=len(page_text)
            ))
            page_num += 1
            current_page = []
            word_count = 0

    # Flush the final partial page
    if current_page:
        page_text = "\n".join(current_page)
        pages.append(PageText(
            page_number=page_num,
            text=page_text,
            char_count=len(page_text)
        ))

    return pages


def extract_text(file_bytes: bytes, filename: str) -> List[PageText]:
    """
    Detect file type from extension and dispatch to the correct extractor.

    Supported: .pdf, .docx, .doc
    Raises ExtractionError for unsupported types.
    """
    ext = (filename or "").rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ExtractionError(
            f"Unsupported file type: .{ext}. Supported formats: pdf, docx"
        )
